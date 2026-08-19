#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CPAHelper 式 Excel -> Word 表格同步：v2 忠实还原（基于操作手册）
作者：Buddy（给俊仪的参考实现 / 事务所 AI 转型）

相比 v1 的关键升级（对应手册发现）：
  1) 绑定锚点 = Excel【名称区域(Named Range)】，前缀 CPAH_，而非原始 A1 区域
     -> 抗行列位移、支持"更换数据源"(同名锚点匹配)
  2) 表头行数：默认只同步数据区，保留 Word 模板表头；可选项同步表头
  3) 列映射：Word 每列 -> Excel 源区域列号（支持长表拆短表）
  4) 同步前排序 + "其他项置底"
  5) 数据源更换：按同名名称区域匹配到新 Excel

依赖：openpyxl, python-docx
用法：
  python excel_word_sync_agent.py --demo               # 生成样例并跑通（含拆分/列映射演示）
  python excel_word_sync_agent.py --sync bindings.json  # 按配置批量同步
"""

import argparse
import json
import os
import re
import shutil
from datetime import datetime


# ============================================================================
# 1. 锚点层：Excel 名称区域（Named Range）
# ============================================================================
def _abs_ref(ref):
    """把 A2:D5 规整成 $A$2:$D$5"""
    out = []
    for p in ref.split(":"):
        m = re.match(r"([A-Za-z]+)(\d+)", p)
        out.append(f"${m.group(1)}${m.group(2)}" if m else p)
    return ":".join(out)


def create_named_range(wb, name, sheet, range_str):
    from openpyxl.workbook.defined_name import DefinedName
    ref = f"{sheet}!{_abs_ref(range_str)}"
    dn = DefinedName(name, attr_text=ref)
    # openpyxl >= 3.1 用字典赋值；老版本用 .add()
    try:
        wb.defined_names[name] = dn
    except AttributeError:
        wb.defined_names.add(dn)


def resolve_named_range(wb, name):
    """返回 (sheet, range_str 去$) 例如 ('费用明细', 'A2:D5')"""
    dn = wb.defined_names[name]
    value = dn.value
    sheet, rng = value.split("!", 1)
    sheet = sheet.strip("'")
    rng = rng.replace("$", "")
    return sheet, rng


def bind(excel_path, sheet, range_str, name=None):
    """建立绑定：复用已有同名区域，否则创建隐藏的 CPAH_ 名称区域。返回名称。"""
    from openpyxl import load_workbook
    wb = load_workbook(excel_path)
    if name is None:
        name = "CPAH_" + re.sub(r"\W", "_", f"{sheet}_{range_str}")
    if name in wb.defined_names:
        # 复用，不覆盖用户原名称
        sheet, range_str = resolve_named_range(wb, name)
    else:
        create_named_range(wb, name, sheet, range_str)
        wb.save(excel_path)
    return name


# ============================================================================
# 2. 读 Excel 区域（按名称区域 or 原始 sheet+range，向后兼容）
# ============================================================================
def read_excel_range(excel_path, sheet, range_str):
    from openpyxl import load_workbook
    wb = load_workbook(excel_path, data_only=True)
    ws = wb[sheet]
    return [[c.value for c in row] for row in ws[range_str]]


def read_by_name(excel_path, name):
    from openpyxl import load_workbook
    wb = load_workbook(excel_path, data_only=True)
    sheet, rng = resolve_named_range(wb, name)
    return read_excel_range(excel_path, sheet, rng)


# ============================================================================
# 3. 预处理：排序 + 其他项置底
# ============================================================================
def _sort_key(v):
    try:
        return (0, float(v))
    except (TypeError, ValueError):
        return (1, "" if v is None else str(v))


def apply_sort(data, cfg):
    if not cfg:
        return data
    col = int(cfg.get("column", 0))
    direction = cfg.get("direction", "asc")
    other_bottom = cfg.get("other_bottom", False)
    other_key = str(cfg.get("other_key", "其他")).strip()

    others = [r for r in data if str(r[col]).strip() == other_key]
    rest = [r for r in data if str(r[col]).strip() != other_key]
    rest.sort(key=lambda r: _sort_key(r[col]), reverse=(direction == "desc"))

    # 注意：只有精确等于"其他"才置底；"其他费用"等正常参与排序（与手册一致）
    return rest + others if other_bottom else others + rest


# ============================================================================
# 4. 预处理：列映射（长表拆短表）
#     Word 列 -> Excel 源区域列号，必须从左到右递增、不重复、不超界
# ============================================================================
def apply_column_map(data, column_map):
    if not column_map:
        return data
    return [[row[i] for i in column_map] for row in data]


def validate_column_map(column_map, n_excel_cols):
    if not column_map:
        return True, ""
    if sorted(column_map) != list(column_map):
        return False, "Excel 列号必须从左到右递增"
    if len(set(column_map)) != len(column_map):
        return False, "不能重复填写同一个 Excel 列"
    if max(column_map) >= n_excel_cols:
        return False, "列号超过名称区域实际列数"
    return True, ""


# ============================================================================
# 5. 写回 Word 表格（原型级：python-docx；生产级见 win32com）
# ============================================================================
def _norm(v):
    """去掉千分符对齐空格带来的前导/尾随空格"""
    if isinstance(v, str):
        return v.strip()
    return v


def write_word_table(docx_path, table_index, data,
                     skip_header_rows=0, backup=True):
    from docx import Document
    if backup:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.copy(docx_path, docx_path + f".bak_{ts}")

    doc = Document(docx_path)
    table = doc.tables[table_index]

    body = [_norm_row(r) for r in data[skip_header_rows:]]
    n_rows = len(body)
    n_excel_cols = max((len(r) for r in body), default=0)
    n_word_cols = len(table.columns)

    while len(table.rows) < n_rows:
        table.add_row()

    for ri, row in enumerate(body):
        for ci in range(min(n_excel_cols, n_word_cols)):
            val = row[ci] if ci < len(row) else ""
            table.cell(ri, ci).text = "" if val is None else str(_norm(val))

    for ri in range(n_rows, len(table.rows)):
        for ci in range(n_word_cols):
            table.cell(ri, ci).text = ""

    doc.save(docx_path)


def _norm_row(row):
    return [_norm(v) for v in row]


# ============================================================================
# 6. 单条绑定同步
# ============================================================================
def sync_binding(b):
    if b.get("name"):
        data = read_by_name(b["excel"], b["name"])
    else:
        data = read_excel_range(b["excel"], b["sheet"], b["range"])

    n_excel_cols = max((len(r) for r in data), default=0)
    ok, msg = validate_column_map(b.get("column_map"), n_excel_cols)
    if not ok:
        print(f"[跳过] {b.get('docx')} 表#{b.get('table_index')} 列映射校验失败：{msg}")
        return

    data = apply_sort(data, b.get("sort"))
    data = apply_column_map(data, b.get("column_map"))
    write_word_table(
        b["docx"], b["table_index"], data,
        b.get("skip_header_rows", 0),
        b.get("backup", True),
    )
    src = b.get("name") or f"{b['sheet']}!{b['range']}"
    print(f"[OK] 已同步：{b['docx']} 表#{b['table_index']} <- {b['excel']} | {src}")


def sync_all(bindings):
    for b in bindings["bindings"]:
        sync_binding(b)


# ============================================================================
# 7. 数据源更换（基于同名锚点匹配）
# ============================================================================
def swap_data_source(bindings, old_excel, new_excel):
    """把绑定里的旧 Excel 路径换成新 Excel（要求新文件含同名 CPAH_ 名称区域）。"""
    from openpyxl import load_workbook
    wb = load_workbook(new_excel)
    names = set(wb.defined_names.keys()) if hasattr(wb.defined_names, "keys") \
        else set(wb.defined_names)
    moved = 0
    for b in bindings["bindings"]:
        if b.get("excel") == old_excel and b.get("name") in names:
            b["excel"] = new_excel
            moved += 1
    print(f"[OK] 更换数据源：{moved} 条绑定已指向 {new_excel}")
    return bindings


# ============================================================================
# 8. 自验证 Demo：含"完整表" + "拆分(列映射)表"
# ============================================================================
def build_demo():
    from openpyxl import Workbook
    from docx import Document

    here = os.path.dirname(os.path.abspath(__file__))
    demo_dir = os.path.join(here, "excel_word_sync_demo")
    os.makedirs(demo_dir, exist_ok=True)

    excel_path = os.path.join(demo_dir, "demo_TB.xlsx")
    docx_path = os.path.join(demo_dir, "demo_报告.docx")
    bpath = os.path.join(demo_dir, "demo_bindings.json")

    # --- Excel 底表（费用明细，区域 A1:D5，含表头）---
    wb = Workbook()
    ws = wb.active
    ws.title = "费用明细"
    ws.append(["科目", "本期", "上期", "变动"])
    ws.append(["其他", 50, 40, 10])
    ws.append(["职工薪酬", 320, 280, 40])
    ws.append(["办公费", 120, 110, 10])
    ws.append(["差旅费", 200, 180, 20])
    wb.save(excel_path)

    # 建立名称区域锚点（含表头 A1:D5）
    name = bind(excel_path, "费用明细", "A1:D5", name="CPAH_费用明细")

    # --- Word 报告：表0=完整4列；表1=拆分后的"科目+变动"2列 ---
    doc = Document()
    doc.add_heading("专项报告 - 费用明细", level=1)
    t0 = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["科目", "本期", "上期", "变动"]):
        t0.rows[0].cells[i].text = h
    t1 = doc.add_table(rows=1, cols=2)   # 拆分表
    for i, h in enumerate(["科目", "变动"]):
        t1.rows[0].cells[i].text = h
    doc.save(docx_path)

    # --- 绑定配置 ---
    bindings = {
        "bindings": [
            {   # 表0：完整同步，表头在名称区域内(skip 1 行)，按本期降序+其他置底
                "docx": docx_path, "table_index": 0,
                "excel": excel_path, "name": name,
                "skip_header_rows": 1,
                "sort": {"column": 1, "direction": "desc",
                         "other_bottom": True, "other_key": "其他"},
                "column_map": None, "backup": True,
            },
            {   # 表1：拆分表，只取 Excel 第 0 列(科目) 与第 3 列(变动)
                "docx": docx_path, "table_index": 1,
                "excel": excel_path, "name": name,
                "skip_header_rows": 1,
                "sort": None,
                "column_map": [0, 3], "backup": True,
            },
        ]
    }
    with open(bpath, "w", encoding="utf-8") as f:
        json.dump(bindings, f, ensure_ascii=False, indent=2)
    return bpath


def print_table(docx_path, table_index):
    from docx import Document
    doc = Document(docx_path)
    t = doc.tables[table_index]
    for r in t.rows:
        print("  | " + " | ".join(str(c.text) for c in r.cells) + " |")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--demo", action="store_true")
    ap.add_argument("--sync", metavar="JSON")
    args = ap.parse_args()

    if args.demo:
        print("== 生成样例 ==")
        bpath = build_demo()
        bindings = json.load(open(bpath, "r", encoding="utf-8"))
        docx = bindings["bindings"][0]["docx"]
        print("同步前 Word 表0(完整)：")
        print_table(docx, 0)
        print("同步前 Word 表1(拆分)：")
        print_table(docx, 1)
        print("\n== 执行同步 ==")
        sync_all(bindings)
        print("\n同步后 Word 表0(完整，按本期降序/其他置底)：")
        print_table(docx, 0)
        print("\n同步后 Word 表1(拆分，仅 科目+变动)：")
        print_table(docx, 1)
        return

    if args.sync:
        sync_all(json.load(open(args.sync, "r", encoding="utf-8")))
        return

    ap.print_help()


if __name__ == "__main__":
    main()
